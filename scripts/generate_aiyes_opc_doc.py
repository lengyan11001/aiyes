from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("docs") / "Aiyes_OPC_API_MCP_Integration_Guide.docx"


def east_asia(run, font="Microsoft YaHei", size=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = east_asia(paragraph.add_run(str(text)), size=9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell_text(tbl.rows[0].cells[idx], header, True)
        shade(tbl.rows[0].cells[idx], "EAF2F8")
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value)
    doc.add_paragraph()
    return tbl


def code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    for line in text.strip("\n").split("\n"):
        run = east_asia(paragraph.add_run(line + "\n"), "Consolas", 8.5)
        run.font.color.rgb = RGBColor(45, 55, 72)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    paragraph._p.get_or_add_pPr().append(shd)


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    east_asia(paragraph.add_run(text))


def main():
    OUT_PATH.parent.mkdir(exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = east_asia(title.add_run("Aiyes OPC / API / MCP 接入说明"), size=22)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = east_asia(meta.add_run("版本：V1.0    日期：2026-05-28"), size=10)
    run.font.color.rgb = RGBColor(89, 99, 110)

    doc.add_paragraph(
        "本文档用于说明外部系统接入 Aiyes 平台的 OPC 数据查询接口、标准 API 接口和 MCP 工具调用方式。"
        "文档仅包含 Aiyes 对外提供的接入信息，不包含模型定价、成本信息或任何非 Aiyes 平台信息。"
    )

    doc.add_heading("1. 接入范围", level=1)
    table(
        doc,
        ["模块", "用途", "入口"],
        [
            ["OPC 查询接口", "用于核验主体、资产、任务、积分/Token 使用、充值订单、发票状态和流水。", "HTTPS REST API"],
            ["标准生成 API", "用于查询模型、提交图片/视频生成任务、查询任务结果。", "HTTPS REST API"],
            ["MCP 接入", "用于通过 MCP 客户端调用 OPC 查询工具。", "Streamable HTTP JSON-RPC"],
        ],
    )
    paragraph = doc.add_paragraph()
    paragraph.add_run("说明：").bold = True
    paragraph.add_run("接口返回的数据以调用方 API Key 所属主体为范围，不返回其他主体数据。")

    doc.add_heading("2. 基础信息", level=1)
    table(
        doc,
        ["项目", "说明"],
        [
            ["平台名称", "Aiyes"],
            ["生产域名", "https://aiyes.vip"],
            ["CMS 域名", "https://cms.aiyes.vip"],
            ["编码格式", "UTF-8"],
            ["请求格式", "application/json"],
            ["时间格式", "ISO 8601，例如 2026-05-01T00:00:00+08:00"],
        ],
    )

    doc.add_heading("3. 鉴权方式", level=1)
    doc.add_paragraph("除公开文档接口外，所有对外接口均使用 Bearer API Key 鉴权。API Key 由 Aiyes 平台账号登录后创建，并仅能访问该账号所属主体的数据。")
    code(
        doc,
        """Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx
Content-Type: application/json""",
    )
    table(
        doc,
        ["状态码", "含义", "处理建议"],
        [
            ["401", "缺少或无效 API Key", "检查 Authorization 头和 API Key 状态。"],
            ["422", "请求参数格式不正确", "根据接口说明修正字段类型或必填项。"],
            ["429", "请求过于频繁", "根据 Retry-After 或业务策略稍后重试。"],
            ["5xx", "服务端异常或临时不可用", "稍后重试，必要时联系 Aiyes 技术支持。"],
        ],
    )

    doc.add_heading("4. OPC 查询接口", level=1)
    doc.add_paragraph("OPC 查询接口用于对接核验场景。所有列表接口支持分页和时间范围过滤。")
    table(
        doc,
        ["参数", "位置", "必填", "说明"],
        [
            ["from", "query", "否", "开始时间，ISO 8601。"],
            ["to", "query", "否", "结束时间，ISO 8601。"],
            ["page", "query", "否", "页码，默认 1。"],
            ["page_size", "query", "否", "每页数量，默认 50，最大 200。"],
        ],
    )
    table(
        doc,
        ["方法", "路径", "说明"],
        [
            ["GET", "/api/v1/opc/docs", "查询接口说明，公开访问。"],
            ["GET", "/api/v1/opc/summary", "查询主体、资产、任务、使用量、充值汇总。"],
            ["GET", "/api/v1/opc/assets", "查询余额、模型、任务数、API Key 数等资产信息。"],
            ["GET", "/api/v1/opc/usage", "查询生成任务和积分/Token 使用记录。"],
            ["GET", "/api/v1/opc/recharges", "查询充值记录。"],
            ["GET", "/api/v1/opc/orders", "查询订单记录，等同充值记录。"],
            ["GET", "/api/v1/opc/invoices", "查询发票状态。未申请开票时返回 not_requested。"],
            ["GET", "/api/v1/opc/ledger", "查询积分流水，包含入账、扣费、调整等记录。"],
        ],
    )

    doc.add_heading("4.1 汇总查询", level=2)
    code(
        doc,
        """curl -X GET "https://aiyes.vip/api/v1/opc/summary?from=2026-05-01T00:00:00+08:00&to=2026-05-31T23:59:59+08:00" \\
  -H "Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx\"""",
    )
    code(
        doc,
        """{
  "subject": {
    "subject_id": "主体ID",
    "subject_name": "主体名称",
    "username": "账号",
    "status": "ACTIVE"
  },
  "period": {
    "from": "2026-05-01T00:00:00.000Z",
    "to": "2026-05-31T15:59:59.000Z"
  },
  "assets": {
    "balance_points": 0,
    "active_api_keys": 1,
    "generation_jobs": 0,
    "completed_jobs": 0,
    "models": []
  },
  "usage": {
    "jobs": 0,
    "completed_jobs": 0,
    "failed_jobs": 0,
    "token_consumed_points": 0
  },
  "recharge": {
    "paid_orders": 0,
    "recharge_points": 0
  }
}""",
    )

    doc.add_heading("4.2 使用记录查询", level=2)
    code(
        doc,
        """curl -X GET "https://aiyes.vip/api/v1/opc/usage?page=1&page_size=50" \\
  -H "Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx\"""",
    )
    table(
        doc,
        ["字段", "说明"],
        [
            ["id / order_no", "平台任务编号，可作为业务核验订单号。"],
            ["upstream_task_id", "任务执行编号，如为空以 id 为准。"],
            ["type", "任务类型：IMAGE 或 VIDEO。"],
            ["model", "调用模型标识。"],
            ["status", "任务状态：PENDING、PROCESSING、COMPLETED、FAILED 等。"],
            ["token_consumed_points", "本次任务消耗的积分/Token 数量。"],
            ["created_at / completed_at", "创建时间和完成时间。"],
        ],
    )

    doc.add_heading("4.3 充值、订单、发票与流水", level=2)
    table(
        doc,
        ["路径", "关键字段", "说明"],
        [
            ["/api/v1/opc/recharges", "order_no、status、credits_points、wx_transaction_id、created_at、paid_at", "充值记录。"],
            ["/api/v1/opc/orders", "order_no、provider、status、credits_points", "订单记录。"],
            ["/api/v1/opc/invoices", "order_no、status、invoice_amount_points_unit、issued_at", "发票状态记录。"],
            ["/api/v1/opc/ledger", "type、amount_points、balance_after_points、job_id、note、created_at", "积分流水。"],
        ],
    )

    doc.add_heading("5. 标准生成 API", level=1)
    doc.add_paragraph("标准生成 API 使用同一套 Bearer API Key。当前仅开放以下模型标识。")
    table(
        doc,
        ["模型标识", "类型", "说明"],
        [
            ["seedance2", "VIDEO", "视频生成模型。"],
            ["openai/gpt-image-2", "IMAGE", "图片生成模型。"],
        ],
    )
    table(
        doc,
        ["方法", "路径", "说明"],
        [
            ["GET", "/api/v1/models", "查询当前可用模型。"],
            ["POST", "/api/v1/images/generations", "提交图片生成任务。"],
            ["POST", "/api/v1/videos/generations", "提交视频生成任务。"],
            ["GET", "/api/v1/jobs/{id}", "查询任务状态和结果。"],
        ],
    )

    doc.add_heading("5.1 一键复制接入模板", level=2)
    doc.add_paragraph(
        "对接方可以直接复制以下模板进行服务端接入。只需要将 AIYES_API_KEY 替换为正式 API Key，"
        "并按业务场景修改 prompt、比例、时长等参数。"
    )
    code(
        doc,
        """# 基础配置
AIYES_BASE_URL="https://aiyes.vip"
AIYES_API_KEY="ak_xxxxxxxxxxxxxxxxxxxxxxxx"

# 1. 查询可用模型
curl -X GET "$AIYES_BASE_URL/api/v1/models" \\
  -H "Authorization: Bearer $AIYES_API_KEY"

# 2. 提交图片生成任务
curl -X POST "$AIYES_BASE_URL/api/v1/images/generations" \\
  -H "Authorization: Bearer $AIYES_API_KEY" \\
  -H "Content-Type: application/json" \\
  -d '{
    "model": "openai/gpt-image-2",
    "prompt": "生成一张商品主视觉图",
    "size": "1K",
    "quality": "medium",
    "async": true
  }'

# 3. 提交视频生成任务
curl -X POST "$AIYES_BASE_URL/api/v1/videos/generations" \\
  -H "Authorization: Bearer $AIYES_API_KEY" \\
  -H "Content-Type: application/json" \\
  -d '{
    "model": "seedance2",
    "prompt": "生成一段商品展示短视频",
    "duration": 4,
    "ratio": "16:9",
    "video_model": "fast",
    "resolution": "720p",
    "async": true
  }'

# 4. 查询任务结果，将 {job_id} 替换为提交任务后返回的 id
curl -X GET "$AIYES_BASE_URL/api/v1/jobs/{job_id}" \\
  -H "Authorization: Bearer $AIYES_API_KEY"

# 5. 查询 OPC 核验汇总
curl -X GET "$AIYES_BASE_URL/api/v1/opc/summary" \\
  -H "Authorization: Bearer $AIYES_API_KEY\"""",
    )

    doc.add_paragraph("Node.js 服务端接入模板：")
    code(
        doc,
        """const AIYES_BASE_URL = "https://aiyes.vip";
const AIYES_API_KEY = process.env.AIYES_API_KEY || "ak_xxxxxxxxxxxxxxxxxxxxxxxx";

async function aiyes(path, options = {}) {
  const res = await fetch(`${AIYES_BASE_URL}${path}`, {
    ...options,
    headers: {
      "Authorization": `Bearer ${AIYES_API_KEY}`,
      "Content-Type": "application/json",
      ...(options.headers || {})
    }
  });
  const data = await res.json();
  if (!res.ok) {
    throw new Error(data?.error?.message || `Aiyes request failed: ${res.status}`);
  }
  return data;
}

async function createVideoTask() {
  return aiyes("/api/v1/videos/generations", {
    method: "POST",
    body: JSON.stringify({
      model: "seedance2",
      prompt: "生成一段商品展示短视频",
      duration: 4,
      ratio: "16:9",
      video_model: "fast",
      resolution: "720p",
      async: true
    })
  });
}

async function queryJob(jobId) {
  return aiyes(`/api/v1/jobs/${jobId}`);
}

async function queryOpcSummary() {
  return aiyes("/api/v1/opc/summary");
}

createVideoTask()
  .then((task) => queryJob(task.id || task.job?.id))
  .then(console.log)
  .catch(console.error);""",
    )

    doc.add_paragraph("MCP 接入模板：")
    code(
        doc,
        """# MCP Endpoint
https://aiyes.vip/api/mcp

# Header
Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx
Content-Type: application/json

# 初始化
{
  "jsonrpc": "2.0",
  "id": 1,
  "method": "initialize",
  "params": {
    "protocolVersion": "2025-06-18",
    "capabilities": {},
    "clientInfo": { "name": "client", "version": "1.0.0" }
  }
}

# 查询工具列表
{
  "jsonrpc": "2.0",
  "id": 2,
  "method": "tools/list"
}

# 调用 OPC 汇总工具
{
  "jsonrpc": "2.0",
  "id": 3,
  "method": "tools/call",
  "params": {
    "name": "opc_summary",
    "arguments": {}
  }
}""",
    )

    doc.add_heading("5.2 图片生成示例", level=2)
    code(
        doc,
        """curl -X POST "https://aiyes.vip/api/v1/images/generations" \\
  -H "Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx" \\
  -H "Content-Type: application/json" \\
  -d '{
    "model": "openai/gpt-image-2",
    "prompt": "生成一张商品主视觉图",
    "size": "1K",
    "quality": "medium",
    "async": true
  }'""",
    )

    doc.add_heading("5.3 视频生成示例", level=2)
    code(
        doc,
        """curl -X POST "https://aiyes.vip/api/v1/videos/generations" \\
  -H "Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx" \\
  -H "Content-Type: application/json" \\
  -d '{
    "model": "seedance2",
    "prompt": "生成一段商品展示短视频",
    "duration": 4,
    "ratio": "16:9",
    "video_model": "fast",
    "resolution": "720p",
    "async": true
  }'""",
    )
    table(
        doc,
        ["字段", "说明"],
        [
            ["model", "固定使用 seedance2。"],
            ["prompt", "生成内容描述。"],
            ["image_url", "可选，参考图 URL。"],
            ["duration", "视频秒数。"],
            ["ratio / aspect_ratio", "画面比例，例如 16:9、9:16、1:1。"],
            ["video_model", "生成模式：fast、standard、fast_vip、standard_vip。"],
            ["resolution", "分辨率，例如 720p、1080p。"],
            ["async", "是否异步返回。建议正式接入使用 true。"],
        ],
    )

    doc.add_heading("5.4 任务查询示例", level=2)
    code(
        doc,
        """curl -X GET "https://aiyes.vip/api/v1/jobs/{job_id}" \\
  -H "Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx\"""",
    )
    table(
        doc,
        ["字段", "说明"],
        [
            ["id", "任务 ID。"],
            ["status", "任务状态。"],
            ["result", "任务结果，完成后返回图片或视频地址等数据。"],
            ["error", "失败原因。"],
            ["chargedCents", "任务消耗积分字段，单位为积分。"],
            ["createdAt / completedAt", "创建时间和完成时间。"],
        ],
    )

    doc.add_heading("6. MCP 接入方式", level=1)
    doc.add_paragraph("Aiyes 提供 MCP Streamable HTTP 接入，适合支持 MCP 的客户端或平台通过工具调用方式查询 OPC 数据。")
    table(
        doc,
        ["项目", "说明"],
        [
            ["MCP Endpoint", "https://aiyes.vip/api/mcp"],
            ["Transport", "Streamable HTTP JSON-RPC over POST"],
            ["鉴权", "Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx"],
        ],
    )
    table(
        doc,
        ["工具名", "说明"],
        [
            ["opc_summary", "查询主体、资产、任务、使用量、充值汇总。"],
            ["opc_assets", "查询主体资产信息。"],
            ["opc_usage", "查询生成任务和积分/Token 使用记录。"],
            ["opc_recharges", "查询充值记录。"],
            ["opc_orders", "查询订单记录。"],
            ["opc_invoices", "查询发票状态。"],
            ["opc_ledger", "查询积分流水。"],
        ],
    )

    doc.add_heading("6.1 MCP 初始化", level=2)
    code(
        doc,
        """POST /api/mcp
Authorization: Bearer ak_xxxxxxxxxxxxxxxxxxxxxxxx
Content-Type: application/json

{
  "jsonrpc": "2.0",
  "id": 1,
  "method": "initialize",
  "params": {
    "protocolVersion": "2025-06-18",
    "capabilities": {},
    "clientInfo": { "name": "client", "version": "1.0.0" }
  }
}""",
    )

    doc.add_heading("6.2 MCP 查询工具列表", level=2)
    code(
        doc,
        """{
  "jsonrpc": "2.0",
  "id": 2,
  "method": "tools/list"
}""",
    )

    doc.add_heading("6.3 MCP 调用工具", level=2)
    code(
        doc,
        """{
  "jsonrpc": "2.0",
  "id": 3,
  "method": "tools/call",
  "params": {
    "name": "opc_summary",
    "arguments": {
      "from": "2026-05-01T00:00:00+08:00",
      "to": "2026-05-31T23:59:59+08:00"
    }
  }
}""",
    )

    doc.add_heading("7. 安全与对接要求", level=1)
    for item in [
        "API Key 仅用于服务端调用，不应写入前端页面、客户端包或公开仓库。",
        "每个接入方建议使用独立 API Key，便于审计、停用和权限管理。",
        "所有接口均使用 HTTPS。",
        "接入方应保存平台返回的任务 ID 或订单号，便于后续核验。",
        "列表接口应按 page 和 page_size 分页拉取，避免一次性拉取过多数据。",
        "本文档不包含任何单价、报价、成本或非 Aiyes 平台接入信息。",
    ]:
        bullet(doc, item)

    doc.add_heading("8. 接入检查清单", level=1)
    table(
        doc,
        ["检查项", "结果"],
        [
            ["已获取生产域名和 API Key", "待确认"],
            ["已完成 /api/v1/opc/summary 调用测试", "待确认"],
            ["已完成分页参数测试", "待确认"],
            ["已完成任务查询或 OPC usage 查询测试", "待确认"],
            ["如使用 MCP，已完成 initialize、tools/list、tools/call 测试", "待确认"],
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = east_asia(footer.add_run("Aiyes 接入说明 - 仅限授权接入使用"), size=9)
    run.font.color.rgb = RGBColor(120, 130, 140)

    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    main()
